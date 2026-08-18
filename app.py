import sys
import os
import re

# 确保项目根目录在 Python 路径中
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from contextlib import asynccontextmanager

from fastapi import FastAPI, HTTPException, UploadFile, File, Request
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, HTMLResponse, StreamingResponse, JSONResponse
from fastapi.exceptions import RequestValidationError
from fastapi.encoders import jsonable_encoder
from pydantic import BaseModel
import uvicorn

import io
import uuid
from datetime import datetime

from config import BASE_DIR
from core import database as core_db
from core.database import db
from core.models import BasicInfo, Education, Internship, Project, Skill, Award, SelfEvaluation, InterviewSession, OtherInfo
from core.deepseek_client import call_deepseek, call_deepseek_json
from services.experience_service import experience_service, sanitize_classification
from services.resume_service import resume_service
from services.jd_service import jd_service
from services.industry_service import industry_service
from services.match_service import compute_match
from services.template_service import template_service
from services.ocr_service import ocr_service
from services.export_service import export_service
from services.cover_letter_service import cover_letter_service
from services.interview_service import interview_service
from services.revise_service import revise_service
from services.template_converter import convert_to_html
from services.company_search_service import analyze_company as search_company_info
from services.dify_client import DIFY_COMPANY_AGENT_API_KEY, DIFY_INTERVIEW_AGENT_API_KEY
from services.diagnosis_service import diagnosis_service
from prompts.experience_parse import EXPERIENCE_PARSE_PROMPT
from prompts.dedup import DEDUP_PROMPT

@asynccontextmanager
async def lifespan(app: FastAPI):
    """应用生命周期：启动时初始化数据库"""
    db.init_db()
    yield


app = FastAPI(title="AI简历定制工具", version="2.2.0", lifespan=lifespan)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://127.0.0.1:8765", "http://localhost:8765"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 静态文件
app.mount("/static", StaticFiles(directory=os.path.join(BASE_DIR, "static")), name="static")


# ==================== 全局异常处理 ====================

@app.exception_handler(RequestValidationError)
async def validation_exception_handler(request: Request, exc: RequestValidationError):
    return JSONResponse(
        status_code=422,
        content={"error": "参数校验失败", "detail": jsonable_encoder(exc.errors())},
    )


@app.exception_handler(HTTPException)
async def http_exception_handler(request: Request, exc: HTTPException):
    return JSONResponse(
        status_code=exc.status_code,
        content={"error": exc.detail, "detail": exc.detail},
    )


@app.exception_handler(Exception)
async def unhandled_exception_handler(request: Request, exc: Exception):
    return JSONResponse(
        status_code=500,
        content={"error": "服务器内部错误", "detail": str(exc)},
    )

@app.get("/")
async def root():
    index_path = os.path.join(BASE_DIR, "static", "index.html")
    if os.path.exists(index_path):
        return FileResponse(index_path)
    return HTMLResponse(
        "<html><body style='font-family:sans-serif;padding:40px;text-align:center'>"
        "<h1>AI简历定制工具</h1><p>服务已启动，等待前端页面...</p></body></html>"
    )

# ==================== Experience Routes ====================

@app.get("/api/experiences/all")
async def get_all_experiences():
    """获取所有经历数据，供前端展示"""
    return {
        "basic_info": experience_service.get_basic_info().to_dict(),
        "education": [e.to_dict() for e in experience_service.list_education()],
        "internships": [i.to_dict() for i in experience_service.list_internships()],
        "projects": [p.to_dict() for p in experience_service.list_projects()],
        "skills": [s.to_dict() for s in experience_service.list_skills()],
        "awards": [a.to_dict() for a in experience_service.list_awards()],
        "others": [o.to_dict() for o in experience_service.list_others()],
        "self_evaluation": experience_service.get_self_evaluation().to_dict(),
    }

class BasicInfoInput(BaseModel):
    name: str = ""
    phone: str = ""
    email: str = ""
    age: str = ""
    job_target: str = ""
    photo_path: str = ""

@app.post("/api/experiences/basic-info")
async def save_basic_info(data: BasicInfoInput):
    info = BasicInfo(**data.model_dump())
    experience_service.save_basic_info(info)
    return {"status": "ok"}

@app.post("/api/experiences/upload-photo")
async def upload_photo(file: UploadFile = File(...)):
    """上传用户照片，自动修正EXIF旋转，返回存储路径"""
    from PIL import Image, ImageOps
    from io import BytesIO

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp'):
        raise HTTPException(400, "仅支持 JPG/PNG/GIF/BMP/WEBP 格式的图片")

    content = await file.read()
    if len(content) > 10 * 1024 * 1024:
        raise HTTPException(400, "图片过大，请上传 10MB 以内的图片")

    # 用Pillow验证文件是有效图片
    try:
        img = Image.open(BytesIO(content))
        img.verify()
    except Exception:
        raise HTTPException(400, "上传文件不是有效的图片")

    # 用Pillow修正EXIF旋转（手机竖拍照片自动转正）
    try:
        img = Image.open(BytesIO(content))
        img = ImageOps.exif_transpose(img)  # 根据EXIF自动旋转
        # 统一转为JPEG保存
        if img.mode in ('RGBA', 'P'):
            img = img.convert('RGB')
        photos_dir = os.path.join(BASE_DIR, "data", "photos")
        os.makedirs(photos_dir, exist_ok=True)
        safe_name = f"photo_{uuid.uuid4().hex}.jpg"
        file_path = os.path.join(photos_dir, safe_name)
        img.save(file_path, 'JPEG', quality=90)
    except Exception:
        raise HTTPException(400, "图片处理失败，请更换图片后重试")

    relative_path = f"data/photos/{safe_name}"
    return {"status": "ok", "photo_path": relative_path}

@app.get("/api/photos/{filename}")
async def get_photo(filename: str):
    """提供照片文件访问，供前端预览（防路径穿越）"""
    photos_dir = os.path.realpath(os.path.join(BASE_DIR, "data", "photos"))
    file_path = os.path.realpath(os.path.join(photos_dir, os.path.basename(filename)))
    if not (file_path == photos_dir or file_path.startswith(photos_dir + os.sep)):
        raise HTTPException(404, "照片不存在")
    if not os.path.exists(file_path):
        raise HTTPException(404, "照片不存在")
    return FileResponse(file_path)

# ======== 具体路由（必须在 /{module} 通配符之前声明） ========

# Self evaluation (必须在 /{module} 之前)
@app.get("/api/experiences/self-evaluation")
async def get_self_eval():
    return experience_service.get_self_evaluation().to_dict()

@app.post("/api/experiences/self-evaluation")
async def save_self_eval(request: Request):
    data = await request.json()
    ev = SelfEvaluation(content=data.get("content", ""))
    experience_service.save_self_evaluation(ev)
    return {"status": "ok"}

# 跨模块移动
class MoveItemInput(BaseModel):
    from_module: str
    item_id: int
    to_module: str

@app.post("/api/experiences/move")
async def move_experience_item(data: MoveItemInput):
    """跨模块移动经历条目（如技能里的证书 → 其他信息）"""
    try:
        new_id = experience_service.move_item(data.from_module, data.item_id, data.to_module)
    except ValueError as e:
        raise HTTPException(400, str(e))
    return {"status": "ok", "id": new_id}


# AI 简历采集提示词（档2：用户复制去豆包等大模型，返回规范 JSON 再粘回工具）
@app.get("/api/experiences/collect-prompt")
async def get_collect_prompt():
    """返回「AI 经历采集提示词」全文，供前端一键复制"""
    from prompts.experience_collect import COLLECT_PROMPT
    return {"prompt": COLLECT_PROMPT}

# AI parse pasted text
class ParseTextInput(BaseModel):
    text: str

@app.post("/api/experiences/parse-text")
async def parse_text(data: ParseTextInput):
    """AI 解析用户粘贴的经历文本为结构化数据"""
    prompt = EXPERIENCE_PARSE_PROMPT.format(user_text=data.text)
    try:
        result = sanitize_classification(await call_deepseek_json(prompt))
        return result
    except Exception:
        raise HTTPException(500, "AI 解析暂不可用，请稍后重试")

# Semantic dedup check
class DedupInput(BaseModel):
    module: str
    new_text: str

class ImportCheckRequest(BaseModel):
    items: dict


@app.post("/api/experiences/clear-all")
async def clear_all_experiences():
    """一键清空所有经历数据（不可恢复，操作前请确认）"""
    tables = ["basic_info", "education", "internships", "projects",
              "skills", "awards", "other_info", "self_evaluation"]
    with db.connection() as conn:
        for t in tables:
            conn.execute(f"DELETE FROM {t}")
    return {"status": "ok", "cleared": tables}


@app.post("/api/experiences/import-check")
async def import_check(data: ImportCheckRequest):
    """AI 导入前去重：本地精确去重（归一化比较），返回重复项列表

    入参 items = parse-text 解析出的完整结构。
    对每个模块的新条目，与已有经历做"去空白/标点/大小写"归一化比较，
    完全相同的标记为重复。零 AI 成本，导入一模一样的经历会被拦下。
    """
    items = data.items or {}
    duplicates = []
    # parse-text 返回 key -> (表名, 去重字段, 展示名称字段)
    mapping = {
        "education": ("education", ["school", "major", "degree"], "school"),
        "work_experience": ("internships", ["company", "position"], "company"),
        "projects": ("projects", ["name"], "name"),
        "skills": ("skills", ["name"], "name"),
        "awards": ("awards", ["name"], "name"),
        "others": ("others", ["title"], "title"),
    }

    def _norm(s):
        if not s:
            return ""
        return re.sub(r"[\s，。、；：:,.!?！？\"\"''\-—|·]+", '', str(s)).lower()

    total = 0
    for api_key, (table, fields, name_field) in mapping.items():
        new_items = items.get(api_key) or []
        if not new_items:
            continue
        existing = getattr(experience_service, f"list_{table}")()
        existing_keys = set()
        for e in existing:
            key = "".join(str(getattr(e, f, "") or "") for f in fields)
            existing_keys.add(_norm(key))
        for idx, ni in enumerate(new_items):
            if not isinstance(ni, dict):
                continue
            total += 1
            key = "".join(str(ni.get(f, "") or "") for f in fields)
            norm_key = _norm(key)
            if norm_key and norm_key in existing_keys:
                duplicates.append({
                    "module": table,
                    "index": idx,
                    "name": str(ni.get(name_field, "") or ""),
                    "reason": "与已有内容完全相同",
                })
    # AI 语义去重（仅 projects/work_experience，一次批量调用；失败降级为只用精确结果）
    suspicious = []
    try:
        suspicious = await _ai_semantic_dedup(items, duplicates)
    except Exception:
        suspicious = []
    return {"duplicates": duplicates, "suspicious": suspicious, "total_checked": total}


async def _ai_semantic_dedup(items: dict, duplicates: list) -> list:
    """对 projects / work_experience 做 AI 语义去重：文字不同但说的是同一件事

    返回 suspicious 列表（疑似重复，供用户二次确认，不自动跳过）。
    """
    result = []
    for api_key, table, fields, name_field in [
        ("projects", "projects", ["name", "actions"], "name"),
        ("work_experience", "internships", ["company", "position"], "company"),
    ]:
        new_items = items.get(api_key) or []
        if not new_items:
            continue
        dup_indexes = {d["index"] for d in duplicates if d["module"] == table}
        remaining = [(i, ni) for i, ni in enumerate(new_items)
                     if isinstance(ni, dict) and i not in dup_indexes]
        if not remaining:
            continue
        existing = getattr(experience_service, f"list_{table}")()
        if not existing:
            continue

        def _line(i, name, extra):
            return f"{i}. {name}：{extra}"

        existing_desc = "\n".join(
            _line(i, str(getattr(e, name_field, "") or ""),
                  "，".join(str(getattr(e, f, "") or "") for f in fields))
            for i, e in enumerate(existing)
        )
        new_desc = "\n".join(
            _line(i, str(ni.get(name_field, "") or ""),
                  "，".join(str(ni.get(f, "") or "") for f in fields))
            for i, ni in remaining
        )
        prompt = f"""你是简历去重助手。判断以下「新条目」是否与「已有条目」说的是同一件事（语义重复，即使文字表达不同）。

模块：{table}
已有条目：
{existing_desc}

新条目：
{new_desc}

请判断每个新条目是否与某个已有条目语义重复（描述的是同一个项目/同一段经历）。
只返回重复项，每个重复项一行，格式严格为：
{{"module": "{table}", "index": 新条目序号, "match": 对应的已有条目名称}}
没有重复就什么都不返回。不要任何解释。"""
        try:
            raw = await call_deepseek(prompt, max_tokens=2048)
            for line in raw.splitlines():
                line = line.strip()
                if not line.startswith("{"):
                    continue
                try:
                    d = json_module.loads(line)
                except Exception:
                    continue
                idx = d.get("index")
                if isinstance(idx, int) and 0 <= idx < len(new_items):
                    result.append({
                        "module": table,
                        "index": idx,
                        "name": str(new_items[idx].get(name_field, "") or ""),
                        "reason": "与已有「%s」语义重复" % d.get("match", "经历"),
                    })
        except Exception:
            continue
    return result


# ==================== 已有简历文件导入（方案B：PDF/Word 抽文本） ====================
@app.post("/api/experiences/import-file")
async def import_experience_file(file: UploadFile = File(...)):
    """上传已有简历文件（PDF / Word .docx / .doc），提取纯文本，供 AI 解析导入

    仅提取文字；图片型 PDF / 扫描件（无文本层）会明确报错，
    引导用户用「方式二：采集提示词」或 OCR（后续升级方案C）。
    """
    import io as _py_io

    filename = file.filename or "resume"
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    content = await file.read()

    try:
        if ext == "pdf":
            from PyPDF2 import PdfReader
            reader = PdfReader(_py_io.BytesIO(content))
            pages = []
            for p in reader.pages:
                t = p.extract_text() or ""
                if t.strip():
                    pages.append(t.strip())
            text = "\n\n".join(pages)
        elif ext == "docx":
            from docx import Document
            doc = Document(_py_io.BytesIO(content))
            parts = [para.text for para in doc.paragraphs if para.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    parts.append(" | ".join([c for c in cells if c]))
            text = "\n".join(parts)
        elif ext == "doc":
            text = await _extract_doc_text(content, filename)
            if text is None:
                raise HTTPException(400, ".doc 老格式解析失败，请用 Word 将文件另存为 .docx 后再上传")
        else:
            raise HTTPException(400, f"不支持的文件格式: .{ext}，请上传 .pdf / .docx / .doc")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"文件解析失败: {str(e)}")

    text = (text or "").strip()
    if not text:
        raise HTTPException(400, "未能从文件中提取到文字：可能是扫描件/图片型 PDF。图片识别暂未开放（见后续升级方案），请改用「采集提示词」方式录入")
    return {"text": text, "filename": filename, "ext": ext}


async def _extract_doc_text(content: bytes, filename: str):
    """.doc 老格式：调本机 Word（win32com）转纯文本；失败返回 None 走友好报错"""
    import os as _os
    import tempfile
    try:
        import win32com.client
        tmp_dir = tempfile.mkdtemp(prefix="resume_doc_")
        src = _os.path.join(tmp_dir, _os.path.basename(filename) or "resume.doc")
        with open(src, "wb") as fh:
            fh.write(content)
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(src, ReadOnly=True)
            text = doc.Content.Text
            doc.Close(False)
            return text
        finally:
            word.Quit()
    except Exception:
        return None


@app.post("/api/experiences/check-duplicate")
async def check_duplicate(data: DedupInput):
    """检查新录入内容是否与已有经历语义重复"""
    existing_items = getattr(experience_service, f"list_{data.module}")()
    if not existing_items:
        return {"is_duplicate": False, "similar_items": [], "suggestion": ""}

    existing_text = "\n".join([
        str(item.to_dict()) for item in existing_items
    ])
    prompt = DEDUP_PROMPT.format(
        module=data.module,
        existing_items=existing_text,
        new_item=data.new_text
    )
    try:
        result = await call_deepseek_json(prompt)
        return result
    except Exception:
        return {"is_duplicate": False, "similar_items": [], "suggestion": "AI 去重检测暂不可用，请稍后重试"}

# ======== 通用 CRUD 端点（/{module} 通配符 — 放在最后） ========

VALID_MODULES = ["education", "internships", "projects", "skills", "awards", "others"]

@app.get("/api/experiences/{module}")
async def list_module(module: str):
    if module not in VALID_MODULES:
        raise HTTPException(404, f"Unknown module: {module}")
    items = getattr(experience_service, f"list_{module}")()
    return [item.to_dict() for item in items]

@app.post("/api/experiences/{module}")
async def add_module_item(module: str, request: Request):
    if module not in VALID_MODULES:
        raise HTTPException(404, f"Unknown module: {module}")
    model_class = {
        "education": Education, "internships": Internship,
        "projects": Project, "skills": Skill, "awards": Award, "others": OtherInfo
    }[module]
    data = await request.json()
    item = model_class(**data)
    item_id = getattr(experience_service, f"add_{module[:-1] if module.endswith('s') else module}")(item)
    return {"status": "ok", "id": item_id}

@app.put("/api/experiences/{module}/{item_id}")
async def update_module_item(module: str, item_id: int, request: Request):
    if module not in VALID_MODULES:
        raise HTTPException(404, f"Unknown module: {module}")
    model_class = {
        "education": Education, "internships": Internship,
        "projects": Project, "skills": Skill, "awards": Award, "others": OtherInfo
    }[module]
    data = await request.json()
    data["id"] = item_id
    item = model_class(**data)
    getattr(experience_service, f"update_{module[:-1] if module.endswith('s') else module}")(item)
    return {"status": "ok"}

@app.delete("/api/experiences/{module}/{item_id}")
async def delete_module_item(module: str, item_id: int):
    if module not in VALID_MODULES:
        raise HTTPException(404, f"Unknown module: {module}")
    getattr(experience_service, f"delete_{module[:-1] if module.endswith('s') else module}")(item_id)
    return {"status": "ok"}

class ReorderInput(BaseModel):
    ids: list[int]

@app.put("/api/experiences/{module}/reorder")
async def reorder_module(module: str, data: ReorderInput):
    if module not in VALID_MODULES:
        raise HTTPException(404, f"Unknown module: {module}")
    experience_service.reorder_items(module, data.ids)
    return {"status": "ok"}

# ==================== Resume Generation Route (MVP) ====================

class GenerateRequest(BaseModel):
    jd_text: str
    template_type: str = "default"
    industry: str = ""  # 用户手动选择的行业，空=自动识别

class MatchPreviewRequest(BaseModel):
    jd_text: str

@app.post("/api/match/preview")
async def match_preview(req: MatchPreviewRequest):
    """生成前匹配度预览 — 计算经历库与JD的关键词命中（不调用AI，零成本）"""
    exp_data = experience_service.export_all()
    experience_text = exp_data["text"]
    if not experience_text.strip():
        return {"score": None, "level": "unknown", "matched": [],
                "missing_keywords": [], "total": 0, "detail": "经历库为空"}
    return compute_match(experience_text, req.jd_text)

@app.post("/api/resumes/generate")
async def generate_resume(req: GenerateRequest):
    """MVP核心接口 — 根据经历库和JD生成简历HTML"""
    # 1. 获取经历和模板
    # 项目级匹配：按 JD 精选最相关的 top_n 个项目（通用逻辑，失败降级为全量）
    selected_ids = None
    try:
        selected_projects = experience_service.select_top_projects(req.jd_text, top_n=2)
        selected_ids = [p.id for p in selected_projects]
    except Exception:
        selected_ids = None
    exp_data = experience_service.export_all(project_ids=selected_ids)
    experience_text = exp_data["text"]
    photo_path = exp_data.get("photo_path", "")

    if not experience_text.strip():
        raise HTTPException(400, "请先在「经历管理」中录入你的经历")

    template_html = template_service.get_template_html(req.template_type)

    # 2. 生成简历（单次AI调用，不再并行生成求职信/面试题/公司分析等）
    resume_result = await resume_service.generate(
        template_html, experience_text, req.jd_text, photo_path,
        industry_override=req.industry,
    )

    if not resume_result.get("html"):
        raise HTTPException(500, "简历生成失败，请重试")

    # 3. 保存记录（生成历史，is_delivered=0 表示仅生成未投递）
    with db.connection() as conn:
        conn.execute(
            """INSERT INTO resume_records
               (jd_text, jd_cleaned, template_name, html_content, created_at, is_delivered)
               VALUES (?, ?, ?, ?, ?, 0)""",
            (req.jd_text, req.jd_text, req.template_type,
             resume_result.get("html", ""),
             datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        )

    return {
        "resume_html": resume_result.get("html"),
        "resume_valid": resume_result.get("valid", False),
        "resume_issues": resume_result.get("issues", []),
        "industry": resume_result.get("industry", ""),
        "match_info": resume_result.get("match_info", {}),
        "selected_projects": [
            {"id": p.id, "name": p.name, "score_tags": (p.tags or "")}
            for p in selected_projects
        ] if selected_projects else [],
    }

# ==================== 简历质量诊断 ====================

class DiagnoseRequest(BaseModel):
    resume_html: str
    jd_text: str = ""

@app.post("/api/resumes/diagnose")
async def diagnose_resume(req: DiagnoseRequest):
    """简历质量诊断 — 客观分（代码算）+ AI找茬 + 边界说明"""
    if not req.resume_html.strip():
        raise HTTPException(400, "resume_html不能为空")
    return await diagnosis_service.diagnose(req.resume_html, req.jd_text)

# ==================== 求职信 & 面试题 ====================

class ExtraRequest(BaseModel):
    jd_text: str

@app.post("/api/resumes/cover-letter")
async def generate_cover_letter(req: ExtraRequest):
    """生成求职信"""
    exp_data = experience_service.export_all()
    if not exp_data["text"].strip():
        raise HTTPException(400, "请先在「经历管理」中录入你的经历")
    result = await cover_letter_service.generate(exp_data["text"], req.jd_text)
    return {"cover_letter": result}

@app.post("/api/resumes/interview-questions")
async def generate_interview_questions(req: ExtraRequest):
    """生成模拟面试题"""
    exp_data = experience_service.export_all()
    if not exp_data["text"].strip():
        raise HTTPException(400, "请先在「经历管理」中录入你的经历")
    result = await interview_service.generate(exp_data["text"], req.jd_text)
    return result

# ==================== 公司分析 ====================

class CompanyAnalyzeRequest(BaseModel):
    company_name: str
    jd_text: str = ""

@app.post("/api/company/analyze")
async def analyze_company(req: CompanyAnalyzeRequest):
    """AI 分析公司，给出求职建议"""
    prompt = f"""你是求职顾问。分析以下公司，给出求职者角度的建议。

公司名称：{req.company_name}
招聘JD：{req.jd_text}

请基于你了解的公开信息分析这家公司。如果没听说过，诚实说数据不足。

然后告诉用户去天眼查查哪些关键数据，**并且教用户怎么判断这些数据是好是坏**。每条 checklist 都要包含：查什么 + 判断标准。

例如：
- "查看注册资本，如果低于100万且实缴资本为0，说明公司资金实力弱，可能是空壳"
- "查看参保人数，如果为0或个位数而JD写着大规模团队，可能虚假宣传"
- "查看成立日期，如果不满1年且无知名投资方，稳定性风险较高"

以JSON返回：
{{
    "summary": "公司基本情况",
    "risk_level": "low/medium/high/unknown",
    "risks": ["风险点"],
    "positives": ["正面因素"],
    "advice": "建议",
    "verdict": "整体靠谱可投 / 有风险建议了解 / 数据不足无法判断",
    "checklist": ["查什么+判断标准", ...]
}}"""
    try:
        result = await call_deepseek_json(prompt)
        return result
    except Exception:
        return {"verdict": "数据不足，无法判断", "risk_level": "unknown", "advice": "AI分析暂不可用，建议通过天眼查等平台查询"}

class CompanyDataInput(BaseModel):
    company_name: str = ""
    raw_data: str  # 用户粘贴的工商数据

@app.post("/api/company/interpret")
async def interpret_company_data(req: CompanyDataInput):
    """用户粘贴天眼查原始数据，AI解读"""
    prompt = f"""你是求职顾问，帮求职者解读公司的工商数据，用通俗语言解释这些数据代表什么、有哪些风险。

公司：{req.company_name}
用户粘贴的工商数据：
{req.raw_data}

请用通俗语言解读以下方面（如果数据中有的话）：
1. 注册资本和实缴资本 — 数字代表什么？实缴低说明什么？
2. 参保人数 — 这个数字说明公司真实规模如何？
3. 成立日期 — 公司处于什么阶段？稳定性如何？
4. 经营异常/行政处罚/法律诉讼 — 严重吗？对求职者有什么影响？
5. 股东/法人变更 — 频繁变更是好是坏？

以JSON返回：
{{
    "analysis": "整体解读（200字内）",
    "risk_signals": ["具体风险信号"],
    "positive_signals": ["积极信号"],
    "verdict": "综合来看靠谱 / 有风险需谨慎 / 数据异常建议避开",
    "action": "给求职者的下一步建议"
}}"""
    try:
        result = await call_deepseek_json(prompt)
        return result
    except Exception:
        return {"verdict": "无法分析", "action": "请确认粘贴的数据格式正确"}

# ==================== Company Search (Agent A: 公司洞察) ====================

class CompanySearchInput(BaseModel):
    company_name: str
    location: str = ""

@app.post("/api/company/search")
async def search_company(req: CompanySearchInput):
    """公司洞察 — 生成6模块结构化分析报告"""
    result = await search_company_info(req.company_name, req.location)
    return result


# ==================== Interview Session (Agent B: 模拟面试) ====================

import uuid as uuid_module
import json as json_module
from core.database import db as _db

@app.post("/api/interview/start")
async def start_interview(request: Request):
    """开始模拟面试 — 生成题目并返回首题"""
    data = await request.json()
    jd_text = data.get("jd_text", "")
    if not jd_text.strip():
        raise HTTPException(400, "请提供目标岗位JD")

    exp_data = experience_service.export_all()
    experience_text = exp_data["text"]
    basic_info = exp_data["text"].split("\n教育背景")[0] if "\n教育背景" in exp_data["text"] else experience_text[:500]

    session_id = f"iv_{uuid_module.uuid4().hex[:12]}"

    # 用 DeepSeek 生成面试题目
    from prompts.interview import build_interview_prompt
    questions_prompt = build_interview_prompt(experience_text, jd_text)
    try:
        questions_data = await call_deepseek_json(questions_prompt)
        all_questions = (
            questions_data.get("tech_questions", []) +
            questions_data.get("project_deep_dive", []) +
            questions_data.get("behavioral_questions", [])
        )
        # 限制最多10题
        if len(all_questions) > 10:
            all_questions = all_questions[:10]
        if not all_questions:
            all_questions = [{"question": "请简单介绍一下你自己和你为什么适合这个岗位？", "purpose": "了解候选人背景"}]
    except Exception:
        all_questions = [{"question": "请简单介绍一下你自己和你为什么适合这个岗位？", "purpose": "了解候选人背景"}]

    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # 存入数据库
    with _db.connection() as conn:
        conn.execute(
            """INSERT INTO interview_sessions
               (session_id, status, basic_info_json, jd_text, experience_text,
                questions_json, current_question_index, chat_history_json, started_at)
               VALUES (?, 'active', ?, ?, ?, ?, 0, '[]', ?)""",
            (session_id, basic_info, jd_text, experience_text,
             json_module.dumps(all_questions, ensure_ascii=False), now)
        )

    first_question = all_questions[0] if all_questions else {"question": "开始面试", "purpose": ""}

    return {
        "session_id": session_id,
        "total_questions": len(all_questions),
        "current_index": 0,
        "question": first_question.get("question", ""),
        "purpose": first_question.get("purpose", ""),
        "started_at": now,
    }


@app.post("/api/interview/answer")
async def submit_answer(request: Request):
    """提交面试回答 — 返回追问或下一题"""
    data = await request.json()
    session_id = data.get("session_id", "")
    answer = data.get("answer", "")

    if not session_id:
        raise HTTPException(400, "缺少 session_id")

    conn = _db.get_connection()
    row = conn.execute(
        "SELECT * FROM interview_sessions WHERE session_id=? AND status='active'",
        (session_id,)
    ).fetchone()
    if not row:
        conn.close()
        raise HTTPException(404, "面试会话不存在或已结束")

    session = InterviewSession.from_row(row)
    questions = json_module.loads(session.questions_json)
    chat_history = json_module.loads(session.chat_history_json)
    current_idx = session.current_question_index

    # 将当前回答加入聊天历史
    current_q = questions[current_idx] if current_idx < len(questions) else {"question": "", "purpose": ""}
    chat_history.append({
        "role": "user",
        "question": current_q.get("question", ""),
        "answer": answer,
        "index": current_idx,
    })

    # AI回应：先点评回答，再决定追问或进入下一题
    is_followup = False
    followup_question = ""
    ai_response = ""

    # 总是让AI回应（不只看长度），对话更自然
    followup_prompt = f"""你正在面试一位候选人。面试氛围是专业但轻松的。

你刚才问的问题：
"{current_q.get('question', '')}"

候选人的回答：
"{answer}"

【你的任务】
1. 先简短自然地对候选人的回答做一个回应——可以是一个简短的肯定、一个相关的追问点、或者自然的过渡语。像真人面试官一样，不要说"收到"、"已记录"这种机械的话。
   - 如果回答有深度：可以说"说得不错，尤其是XX部分"、"这个经验很有价值"
   - 如果回答有模糊点：自然追问那个点
   - 如果需要进入下一题：用过渡语如"了解了，那我们换个话题..."

2. 然后决定下一步：
   - 如果回答有值得深挖但不清晰的地方 → 自然追问，深入了解
   - 如果回答已经足够完整 → 说"__NEXT__"进入下一题

【输出格式】
先写你的回应（1-2句话的自然对话），再另起一行写：
__FOLLOWUP__（如果要追问的话）
或
__NEXT__（如果进入下一题）

示例1（追问）：
你在微服务拆分上的经验确实很丰富。能具体说说当时是怎么决定按业务域还是按技术层拆分的吗？
__FOLLOWUP__

示例2（下一题）：
了解了，这个项目经历很有说服力。那我们换个方向聊聊...
__NEXT__"""

    ai_text = await call_deepseek(followup_prompt, max_tokens=256)
    if ai_text:
        # 解析AI回应
        if "__FOLLOWUP__" in ai_text:
            is_followup = True
            # 提取回应和追问（分隔符之后才是真正的追问）
            parts = ai_text.split("__FOLLOWUP__", 1)
            ai_response = parts[0].strip()
            followup_question = parts[1].strip() if len(parts) > 1 else ""
            if not followup_question:
                followup_question = ai_response
            # AI 回应 + 追问一并写入聊天历史
            chat_history.append({
                "role": "ai_response",
                "question": current_q.get("question", ""),
                "answer": answer,
                "text": ai_response,
                "followup": followup_question,
                "index": current_idx,
            })
        elif "__NEXT__" in ai_text:
            parts = ai_text.split("__NEXT__", 1)
            ai_response = parts[0].strip()
            # 将AI回应追加到聊天记录
            if ai_response:
                chat_history.append({
                    "role": "ai_response",
                    "question": current_q.get("question", ""),
                    "answer": answer,
                    "text": ai_response,
                    "followup": "",
                    "index": current_idx,
                })
        else:
            # 没有标记时，整个文本作为追问
            is_followup = True
            followup_question = ai_text.strip()
            chat_history.append({
                "role": "ai_response",
                "question": current_q.get("question", ""),
                "answer": answer,
                "text": "",
                "followup": followup_question,
                "index": current_idx,
            })

    if not is_followup:
        # 进入下一题
        current_idx += 1

    is_complete = current_idx >= len(questions)

    # 更新数据库
    conn.execute(
        """UPDATE interview_sessions
           SET current_question_index=?, chat_history_json=?
           WHERE session_id=?""",
        (current_idx, json_module.dumps(chat_history, ensure_ascii=False), session_id)
    )
    conn.commit()

    if is_complete:
        # 生成评估报告
        evaluation = await _generate_evaluation(chat_history, session.jd_text, session.experience_text)
        conn.execute(
            """UPDATE interview_sessions
               SET status='completed', evaluation_json=?, ended_at=?
               WHERE session_id=?""",
            (json_module.dumps(evaluation, ensure_ascii=False),
             datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
             session_id)
        )
        conn.commit()
        conn.close()
        return {
            "session_id": session_id,
            "is_complete": True,
            "evaluation": evaluation,
        }

    conn.close()

    next_q = questions[current_idx] if current_idx < len(questions) else {"question": "", "purpose": ""}
    return {
        "session_id": session_id,
        "is_complete": False,
        "is_followup": is_followup,
        "current_index": current_idx,
        "total_questions": len(questions),
        "question": followup_question if is_followup else next_q.get("question", ""),
        "purpose": "" if is_followup else next_q.get("purpose", ""),
        "ai_text": ai_response,
    }


@app.post("/api/interview/end")
async def end_interview(request: Request):
    """主动结束面试 — 生成评估"""
    data = await request.json()
    session_id = data.get("session_id", "")

    with _db.connection() as conn:
        row = conn.execute(
            "SELECT * FROM interview_sessions WHERE session_id=? AND status='active'",
            (session_id,)
        ).fetchone()
    if not row:
        raise HTTPException(404, "面试会话不存在或已结束")

    session = InterviewSession.from_row(row)
    chat_history = json_module.loads(session.chat_history_json)

    evaluation = await _generate_evaluation(chat_history, session.jd_text, session.experience_text)

    with _db.connection() as conn:
        conn.execute(
            """UPDATE interview_sessions
               SET status='completed', evaluation_json=?, ended_at=?
               WHERE session_id=?""",
            (json_module.dumps(evaluation, ensure_ascii=False),
             datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
             session_id)
        )

    return {
        "session_id": session_id,
        "is_complete": True,
        "evaluation": evaluation,
    }


async def _generate_evaluation(chat_history: list, jd_text: str, experience_text: str) -> dict:
    """生成面试评估报告"""
    qa_text = ""
    for item in chat_history:
        if item.get("role") == "ai_response":
            # AI 回应条目：展示回应/追问内容，避免遗漏
            ai_text = item.get("text", "") or item.get("followup", "")
            if ai_text:
                qa_text += f"AI: {ai_text}\n\n"
            continue
        qa_text += f"Q: {item.get('question', '')}\nA: {item.get('answer', '')}\n\n"

    prompt = f"""你是面试评估专家。基于以下面试对话，给出结构化评估。

【岗位JD】
{jd_text}

【候选人背景】
{experience_text[:1000]}

【面试对话记录】
{qa_text}

以JSON返回评估报告：
{{
    "overall_match": "85（示例：匹配度百分比）",
    "overall_comment": "整体简评",
    "dimensions": [
        {{"name": "经历匹配度", "stars": 4, "comment": "评分理由"}},
        {{"name": "逻辑与表达", "stars": 4, "comment": "评分理由"}},
        {{"name": "技术深度", "stars": 3, "comment": "评分理由"}}
    ],
    "highlights": ["亮点1", "亮点2"],
    "improvements": ["改进点1", "改进点2"],
    "preparation_advice": "具体的学习或准备建议"
}}

要求：评分1-5星，建议要具体可操作。"""

    try:
        return await call_deepseek_json(prompt)
    except Exception:
        return {
            "overall_match": "N/A",
            "overall_comment": "评估生成失败，请重试",
            "dimensions": [],
            "highlights": [],
            "improvements": [],
            "preparation_advice": "",
        }


# ==================== Delivery Records (Agent C: 投递助手) ====================

class DeliverySubmitRequest(BaseModel):
    resume_html: str
    jd_text: str
    company_name: str = ""
    job_title: str = ""
    delivery_url: str = ""

@app.post("/api/delivery/submit")
async def submit_delivery(req: DeliverySubmitRequest):
    """投递此岗位 — 写入投递记录"""
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # 从JD中提取公司名和岗位名（如未提供）
    company_name = req.company_name
    job_title = req.job_title
    if not company_name or not job_title:
        try:
            jd_result = await jd_service.clean(req.jd_text)
            if not company_name:
                company_name = jd_result.get("company_name", "")
            if not job_title:
                job_title = jd_result.get("job_title", "")
        except Exception:
            pass

    # 检查重复投递
    with _db.connection() as conn:
        existing = conn.execute(
            """SELECT id FROM resume_records
               WHERE jd_text=? AND is_delivered=1
               ORDER BY created_at DESC LIMIT 1""",
            (req.jd_text,)
        ).fetchone()

    if existing:
        return {"success": False, "error": "该JD已投递过，请勿重复投递", "duplicate_id": existing["id"]}

    # 写入投递记录
    with _db.connection() as conn:
        conn.execute(
            """INSERT INTO resume_records
               (jd_text, html_content, company_name, job_title, is_delivered,
                delivery_time, delivery_url, delivery_status, created_at)
               VALUES (?, ?, ?, ?, 1, ?, ?, 'delivered', ?)""",
            (req.jd_text, req.resume_html, company_name, job_title,
             now, req.delivery_url, now)
        )
        record_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]

    return {
        "success": True,
        "record_id": record_id,
        "company_name": company_name,
        "job_title": job_title,
        "delivery_time": now,
        "message": "简历已复制到剪贴板",
    }


@app.get("/api/delivery/records")
async def list_delivery_records(search: str = "", page: int = 1, page_size: int = 20):
    """投递记录列表 — 支持搜索和分页"""
    offset = (page - 1) * page_size
    with _db.connection() as conn:
        if search:
            rows = conn.execute(
                """SELECT id, company_name, job_title, delivery_time, delivery_status, created_at
                   FROM resume_records
                   WHERE is_delivered=1
                     AND (company_name LIKE ? OR job_title LIKE ?)
                   ORDER BY delivery_time DESC
                   LIMIT ? OFFSET ?""",
                (f"%{search}%", f"%{search}%", page_size, offset)
            ).fetchall()

            total_row = conn.execute(
                """SELECT COUNT(*) as cnt FROM resume_records
                   WHERE is_delivered=1
                     AND (company_name LIKE ? OR job_title LIKE ?)""",
                (f"%{search}%", f"%{search}%")
            ).fetchone()
        else:
            rows = conn.execute(
                """SELECT id, company_name, job_title, delivery_time, delivery_status, created_at
                   FROM resume_records
                   WHERE is_delivered=1
                   ORDER BY delivery_time DESC
                   LIMIT ? OFFSET ?""",
                (page_size, offset)
            ).fetchall()

            total_row = conn.execute(
                "SELECT COUNT(*) as cnt FROM resume_records WHERE is_delivered=1"
            ).fetchone()

    total = total_row["cnt"] if total_row else 0

    records = []
    for row in rows:
        records.append({
            "id": row["id"],
            "company_name": row["company_name"],
            "job_title": row["job_title"],
            "delivery_time": row["delivery_time"],
            "delivery_status": row["delivery_status"],
            "created_at": row["created_at"],
        })

    return {
        "records": records,
        "total": total,
        "page": page,
        "page_size": page_size,
    }


@app.get("/api/delivery/records/{record_id}")
async def get_delivery_detail(record_id: int):
    """投递详情 — 含历史简历HTML"""
    with _db.connection() as conn:
        row = conn.execute(
            "SELECT * FROM resume_records WHERE id=? AND is_delivered=1",
            (record_id,)
        ).fetchone()

    if not row:
        raise HTTPException(404, "投递记录不存在")

    return {
        "id": row["id"],
        "company_name": row["company_name"],
        "job_title": row["job_title"],
        "jd_text": row["jd_text"],
        "html_content": row["html_content"],
        "delivery_time": row["delivery_time"],
        "delivery_url": row["delivery_url"],
        "delivery_status": row["delivery_status"],
        "created_at": row["created_at"],
    }


@app.delete("/api/delivery/records/{record_id}")
async def delete_delivery_record(record_id: int):
    """删除投递记录"""
    with _db.connection() as conn:
        conn.execute("DELETE FROM resume_records WHERE id=? AND is_delivered=1", (record_id,))
    return {"status": "ok"}


# ==================== Template Routes ====================

@app.get("/api/templates")
async def list_templates():
    return template_service.list_templates()

class TemplateUploadInput(BaseModel):
    name: str
    html_content: str

@app.post("/api/templates/upload")
async def upload_template(data: TemplateUploadInput):
    result = await template_service.upload_template(data.name, data.html_content)
    return result


@app.post("/api/templates/import-file")
async def import_template_file(file: UploadFile = File(...)):
    """导入模板文件 — 支持 HTML/Word/PDF，自动转换为 HTML 后保存"""
    import aiofiles

    filename = file.filename or "template"
    ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''

    # 读取文件内容
    content = await file.read()

    if ext in ('html', 'htm'):
        # 直接使用 HTML 内容
        try:
            html_content = content.decode('utf-8')
        except UnicodeDecodeError:
            html_content = content.decode('gbk', errors='replace')
        name = filename.rsplit('.', 1)[0]
    elif ext in ('docx', 'doc', 'pdf'):
        # 转换 Word/PDF 为 HTML
        try:
            html_content = convert_to_html(content, filename)
        except RuntimeError as e:
            raise HTTPException(500, str(e))
        except Exception as e:
            raise HTTPException(500, f"文件转换失败: {str(e)}")
        name = filename.rsplit('.', 1)[0]
    else:
        raise HTTPException(400, f"不支持的文件格式: .{ext}，支持 .html / .docx / .pdf")

    if not html_content or not html_content.strip():
        raise HTTPException(400, "转换后的模板内容为空")

    result = await template_service.upload_template(name, html_content)
    return result

@app.delete("/api/templates/{template_id}")
async def delete_template(template_id: int):
    template_service.delete_template(template_id)
    return {"status": "ok"}

# ==================== JD Routes ====================

@app.post("/api/jd/clean")
async def clean_jd(request: Request):
    data = await request.json()
    jd_text = data.get("jd_text", "")
    if not jd_text:
        raise HTTPException(400, "jd_text is required")
    result = await jd_service.clean(jd_text)
    return result

class IndustryAnalysisRequest(BaseModel):
    jd_text: str
    industry: str = ""  # 用户手动指定行业（可选，空=自动识别）

@app.post("/api/jd/industry-analysis")
async def industry_analysis(req: IndustryAnalysisRequest):
    """行业侧重点分析 — 识别行业并给出该行业简历侧重点"""
    if not req.jd_text.strip():
        raise HTTPException(400, "JD不能为空")
    return await industry_service.analyze(req.jd_text, industry_override=req.industry)

# ==================== OCR Routes ====================

@app.post("/api/ocr/extract")
async def ocr_extract(files: list[UploadFile] = File(...)):
    """上传截图，返回 OCR 识别文本"""
    import aiofiles

    texts = []
    for file in files:
        tmp_path = os.path.join(BASE_DIR, "data", f"ocr_{uuid.uuid4().hex}.png")
        async with aiofiles.open(tmp_path, 'wb') as f:
            content = await file.read()
            await f.write(content)

        text = await ocr_service._ocr_single(tmp_path)
        texts.append({"filename": file.filename, "text": text})

        try:
            os.remove(tmp_path)
        except Exception:
            pass

    merged = "\n\n---\n\n".join([t["text"] for t in texts])
    return {"texts": texts, "merged_text": merged}

# ==================== Export Routes ====================

@app.post("/api/export/pdf")
async def export_pdf(request: Request):
    """导出 PDF — 优先使用 xhtml2pdf，失败时降级为 HTML 下载"""
    data = await request.json()
    html_content = data.get("html_content", "")
    if not html_content:
        raise HTTPException(400, "html_content is required")

    try:
        pdf_bytes = export_service.to_pdf(html_content)
        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=resume.pdf"}
        )
    except (OSError, RuntimeError) as e:
        # xhtml2pdf 不可用时（中文/复杂 CSS 支持有限），降级为 HTML 文件下载
        # 用户可以用浏览器的 Ctrl+P 打印为 PDF
        html_full = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>简历</title>
<style>@media print {{ body {{ -webkit-print-color-adjust: exact; }} }}</style>
</head><body>{html_content}</body></html>"""
        return StreamingResponse(
            io.BytesIO(html_full.encode("utf-8")),
            media_type="text/html",
            headers={"Content-Disposition": "attachment; filename=resume.html"}
        )

# ==================== AI Chat Routes ====================

class ModifyRequest(BaseModel):
    selected_text: str
    instruction: str
    experience_context: str = ""

@app.post("/api/ai/modify")
async def ai_modify(data: ModifyRequest):
    """选中文本 + 修改要求 → AI 返回修改后的文本段落"""
    prompt = f"""你是简历修改助手。用户选中了简历中的一段文字，要求修改。

【选中的原文】
{data.selected_text}

【用户的修改要求】
{data.instruction}

【用户的经历库参考（如有）】
{data.experience_context}

请只返回修改后的文本段落（不要包含任何解释、不要返回整份简历、不要加markdown代码块）。
直接输出替换选中段落后应该写入的新文本。保持相似的篇幅和格式。"""
    result = await call_deepseek(prompt, max_tokens=1024)
    return {"modified_text": result}

# ==================== AI Revise Routes ====================

class ReviseRequest(BaseModel):
    current_html: str
    instruction: str

@app.post("/api/resumes/revise")
async def revise_resume(data: ReviseRequest):
    """AI修改简历 - 完整HTML diff模式，CSS/图片格式保护"""
    revised_html, message = await revise_service.revise(
        data.current_html, data.instruction
    )
    if revised_html is None:
        raise HTTPException(500, message)
    return {"success": True, "revised_html": revised_html, "message": message}


class AutoTrimRequest(BaseModel):
    resume_html: str
    jd_text: str = ""


@app.post("/api/resumes/auto-expand")
async def auto_expand_resume(data: AutoTrimRequest):
    """AI 自动扩写：内容不足一页时基于真实经历扩充（不编造事实）"""
    expanded_html, message = await revise_service.expand(data.resume_html, data.jd_text)
    if expanded_html is None:
        raise HTTPException(500, message)
    return {"success": True, "revised_html": expanded_html, "message": message}


@app.post("/api/resumes/auto-trim")
async def auto_trim_resume(data: AutoTrimRequest):
    """AI 自动精简简历以适配一页（直接输出干净 HTML，无 diff 标记）"""
    trimmed_html, message = await revise_service.trim(data.resume_html, data.jd_text)
    if trimmed_html is None:
        raise HTTPException(500, message)
    return {"success": True, "revised_html": trimmed_html, "message": message}


class AcceptRevisionRequest(BaseModel):
    html_content: str

@app.post("/api/resumes/accept-revision")
async def accept_revision(data: AcceptRevisionRequest):
    """接受修改 - 清除del/ins标记，返回干净的HTML"""
    clean_html = revise_service.accept(data.html_content)
    if clean_html is None:
        raise HTTPException(400, "HTML内容为空")
    return {"success": True, "clean_html": clean_html}


if __name__ == "__main__":
    # reload 仅开发调试用（设置环境变量 RESUME_DEV=1 开启）；
    # 默认单进程：关掉服务窗口即彻底停止，避免残留 worker 占端口导致下次启动失败
    uvicorn.run("app:app", host="127.0.0.1", port=8765, reload=os.environ.get("RESUME_DEV") == "1")
