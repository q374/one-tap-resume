"""模板格式转换：Word (.docx) / PDF → HTML（保留颜色、图形、排版）"""

import io
import re
import base64
import os


# ─── Word (.docx) 转换 ──────────────────────────────────

def _rgb_to_hex(rgb) -> str | None:
    """将 python-docx RGBColor 或字符串转为 #RRGGBB"""
    if rgb is None:
        return None
    try:
        if hasattr(rgb, 'hex'):
            return f'#{rgb.hex}'
        s = str(rgb)
        if s.startswith('#'):
            return s
        return f'#{s}'
    except Exception:
        return None


def _extract_run_style(run) -> str:
    """从一个 Run 中提取所有内联样式，返回 CSS 字符串"""
    styles = []
    font = run.font

    if font.bold:
        styles.append('font-weight:bold')
    if font.italic:
        styles.append('font-style:italic')
    if font.underline:
        styles.append('text-decoration:underline')
    if font.size:
        pt = font.size.pt
        styles.append(f'font-size:{pt}pt')

    color = font.color
    if color and color.rgb:
        hex_c = _rgb_to_hex(color.rgb)
        if hex_c and hex_c.lower() != '#000000':
            styles.append(f'color:{hex_c}')

    if font.name:
        styles.append(f"font-family:'{font.name}'")

    # 高亮
    try:
        from docx.oxml.ns import qn
        rpr = run._element.find(qn('w:rPr'))
        if rpr is not None:
            highlight = rpr.find(qn('w:highlight'))
            if highlight is not None:
                val = highlight.get(qn('w:val'))
                highlight_map = {
                    'yellow': '#ffff00', 'green': '#00ff00', 'cyan': '#00ffff',
                    'magenta': '#ff00ff', 'blue': '#0000ff', 'red': '#ff0000',
                    'darkBlue': '#00008b', 'darkCyan': '#008b8b', 'darkGreen': '#006400',
                    'darkMagenta': '#8b008b', 'darkRed': '#8b0000', 'darkYellow': '#808000',
                    'darkGray': '#808080', 'lightGray': '#d3d3d3', 'black': '#000000',
                }
                if val and val != 'none':
                    bg = highlight_map.get(val, f'#{val}')
                    styles.append(f'background-color:{bg}')
    except Exception:
        pass

    return ';'.join(styles) if styles else ''


def _extract_para_style(para) -> str:
    """提取段落的块级样式"""
    styles = []
    pf = para.paragraph_format

    if pf.alignment:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        align_map = {
            WD_ALIGN_PARAGRAPH.CENTER: 'text-align:center',
            WD_ALIGN_PARAGRAPH.RIGHT: 'text-align:right',
            WD_ALIGN_PARAGRAPH.JUSTIFY: 'text-align:justify',
        }
        if pf.alignment in align_map:
            styles.append(align_map[pf.alignment])

    # 段落缩进
    if pf.first_line_indent:
        styles.append(f'text-indent:{pf.first_line_indent.pt}pt')
    if pf.left_indent:
        styles.append(f'padding-left:{pf.left_indent.pt}pt')

    # 段落背景色 (shading)
    try:
        from docx.oxml.ns import qn
        ppr = para._element.find(qn('w:pPr'))
        if ppr is not None:
            shd = ppr.find(qn('w:shd'))
            if shd is not None:
                fill = shd.get(qn('w:fill'))
                if fill and fill != 'auto' and fill != '000000':
                    styles.append(f'background-color:#{fill}')
    except Exception:
        pass

    return ';'.join(styles) if styles else ''


def _extract_images_from_docx(doc) -> dict[str, str]:
    """从 docx 中提取所有图片，返回 {rId: base64_data_uri}"""
    images = {}
    try:
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    img_bytes = rel.target_part.blob
                    ext = os.path.splitext(rel.target_part.partname)[1].lower()
                    mime_map = {'.png': 'image/png', '.jpg': 'image/jpeg',
                                '.jpeg': 'image/jpeg', '.gif': 'image/gif',
                                '.bmp': 'image/bmp', '.webp': 'image/webp'}
                    mime = mime_map.get(ext, 'image/png')
                    b64 = base64.b64encode(img_bytes).decode('utf-8')
                    images[rel.rId] = f'data:{mime};base64,{b64}'
                except Exception:
                    continue
    except Exception:
        pass
    return images


def _convert_paragraph(para, images: dict[str, str]) -> str:
    """将单个段落转为 HTML"""
    # 检测段落中的图片
    try:
        from docx.oxml.ns import qn
        drawings = para._element.findall('.//' + qn('wp:inline'))
        drawings.extend(para._element.findall('.//' + qn('wp:anchor')))
        img_tags = []
        for drawing in drawings:
            blips = drawing.findall('.//' + qn('a:blip'))
            for blip in blips:
                embed = blip.get(qn('r:embed'))
                if embed and embed in images:
                    # 获取图片尺寸
                    extents = drawing.findall('.//' + qn('wp:extent'))
                    w, h = '', ''
                    for ext in extents:
                        cx = ext.get('cx')
                        cy = ext.get('cy')
                        if cx:
                            w = f'width:{int(cx)/914400:.2f}in'
                        if cy:
                            h = f'height:{int(cy)/914400:.2f}in'
                    size_style = ';'.join(filter(None, [w, h]))
                    img_tags.append(
                        f'<img src="{images[embed]}" style="{size_style};max-width:100%">'
                    )
        if img_tags:
            return ''.join(img_tags)
    except Exception:
        pass

    # 文字段落
    para_style = _extract_para_style(para)
    has_content = False
    parts = []

    for run in para.runs:
        text = run.text
        if not text:
            continue
        has_content = True
        run_style = _extract_run_style(run)

        if run_style:
            parts.append(f'<span style="{run_style}">{text}</span>')
        else:
            parts.append(text)

    if not has_content:
        return '<p>&nbsp;</p>'

    content = ''.join(parts)
    if para_style:
        return f'<p style="{para_style}">{content}</p>'
    return f'<p>{content}</p>'


def _convert_table(table, images: dict[str, str]) -> str:
    """将 Word 表格转为 HTML 表格，保留单元格样式"""
    rows_html = []
    for row in table.rows:
        cells_html = []
        for cell in row.cells:
            # 提取单元格背景色
            cell_style = ''
            try:
                from docx.oxml.ns import qn
                tcpr = cell._element.find(qn('w:tcPr'))
                if tcpr is not None:
                    shd = tcpr.find(qn('w:shd'))
                    if shd is not None:
                        fill = shd.get(qn('w:fill'))
                        if fill and fill != 'auto' and fill != '000000':
                            cell_style = f'style="background-color:#{fill}"'
            except Exception:
                pass

            cell_text = []
            for para in cell.paragraphs:
                text = para.text.strip()
                if text:
                    cell_text.append(text)
            cells_html.append(f'<td {cell_style}>{" ".join(cell_text)}</td>')

        rows_html.append('<tr>' + ''.join(cells_html) + '</tr>')

    return '<table style="width:100%;border-collapse:collapse;margin:8px 0">' + '\n'.join(rows_html) + '</table>'


def convert_docx_to_html(file_bytes: bytes, filename: str) -> str:
    """将 .docx 文件转换为 HTML

    优先使用 Microsoft Word COM 引擎（完美保留格式），
    不可用时降级为 python-docx 提取。
    """
    # ── 方案A: Word COM 引擎（最佳） ──
    html = _try_word_com_conversion(file_bytes, filename)
    if html:
        return html

    # ── 方案B: python-docx 提取（降级） ──
    return _convert_with_python_docx(file_bytes, filename)


def _try_word_com_conversion(file_bytes: bytes, filename: str) -> str | None:
    """尝试用 Word COM 将 docx 转为 HTML，失败返回 None"""
    import tempfile, os
    tmp_docx = None
    tmp_html = None
    try:
        import win32com.client
        # 写临时 docx 文件
        fd_docx, tmp_docx = tempfile.mkstemp(suffix='.docx')
        os.close(fd_docx)
        with open(tmp_docx, 'wb') as f:
            f.write(file_bytes)

        tmp_html = tmp_docx.replace('.docx', '.html')

        word = win32com.client.Dispatch('Word.Application')
        word.Visible = False
        try:
            doc = word.Documents.Open(tmp_docx)
            doc.SaveAs(tmp_html, FileFormat=8)  # wdFormatHTML
            doc.Close()
        finally:
            word.Quit()

        # 读取 HTML（Word 导出用 GBK 编码）
        for enc in ['gbk', 'gb18030', 'utf-8', 'gb2312']:
            try:
                with open(tmp_html, 'r', encoding=enc) as f:
                    return f.read()
            except (UnicodeDecodeError, LookupError):
                continue
        return None
    except Exception:
        return None
    finally:
        for p in [tmp_docx, tmp_html]:
            if p and os.path.exists(p):
                try:
                    os.remove(p)
                except Exception:
                    pass


def _convert_with_python_docx(file_bytes: bytes, filename: str) -> str:
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("需要安装 python-docx: pip install python-docx")

    doc = Document(io.BytesIO(file_bytes))
    images = _extract_images_from_docx(doc)
    name = filename.replace('.docx', '').replace('.doc', '')

    # 提取页面背景色
    page_bg = ''
    try:
        from docx.oxml.ns import qn
        sect_prs = doc.element.findall('.//' + qn('w:sectPr'))
        for sect_pr in sect_prs:
            bg_elem = sect_pr.find(qn('w:background'))
            if bg_elem is not None:
                bg_color = bg_elem.get(qn('w:color'))
                if bg_color and bg_color != 'auto':
                    page_bg = f'background-color:#{bg_color};'
                    break
    except Exception:
        pass

    body_style = f' style="{page_bg}"' if page_bg else ''

    parts = []
    parts.append('<!DOCTYPE html>')
    parts.append('<html lang="zh-CN">')
    parts.append('<head><meta charset="UTF-8">')
    parts.append(f'<title>{name}</title>')
    parts.append('<style>')
    parts.append(f'body{{margin:0;padding:0;{page_bg}}}')
    parts.append('table{border-collapse:collapse;}')
    parts.append('td,th{border:1px solid #ccc;padding:4px 8px;}')
    parts.append('</style>')
    parts.append(f'</head><body{body_style}>')

    # 逐元素处理（段落 + 表格）
    body = doc.element.body
    try:
        from docx.oxml.ns import qn
        for child in body:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

            if tag == 'p':
                # 找到对应的 Paragraph 对象
                para = None
                for p in doc.paragraphs:
                    if p._element is child:
                        para = p
                        break
                if para is not None:
                    # 检查段落样式是否为标题
                    style_name = para.style.name if para.style else ''
                    heading_level = None
                    if 'Heading 1' in style_name or 'heading 1' in style_name.lower():
                        heading_level = 'h1'
                    elif 'Heading 2' in style_name or 'heading 2' in style_name.lower():
                        heading_level = 'h2'
                    elif 'Heading 3' in style_name or 'heading 3' in style_name.lower():
                        heading_level = 'h3'

                    if heading_level:
                        text = para.text.strip()
                        para_style = _extract_para_style(para)
                        if para_style:
                            parts.append(f'<{heading_level} style="{para_style}">{text}</{heading_level}>')
                        else:
                            parts.append(f'<{heading_level}>{text}</{heading_level}>')
                    else:
                        parts.append(_convert_paragraph(para, images))

            elif tag == 'tbl':
                # 找到对应的 Table 对象
                tbl = None
                for t in doc.tables:
                    if t._element is child:
                        tbl = t
                        break
                if tbl is not None:
                    parts.append(_convert_table(tbl, images))
    except Exception:
        # 退化：只用段落顺序处理
        body_elements = []
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ''
            if 'Heading 1' in style_name or 'heading 1' in style_name.lower():
                body_elements.append(f'<h1>{para.text.strip()}</h1>')
            elif 'Heading 2' in style_name or 'heading 2' in style_name.lower():
                body_elements.append(f'<h2>{para.text.strip()}</h2>')
            elif 'Heading 3' in style_name or 'heading 3' in style_name.lower():
                body_elements.append(f'<h3>{para.text.strip()}</h3>')
            else:
                body_elements.append(_convert_paragraph(para, images))

        for table in doc.tables:
            body_elements.append(_convert_table(table, images))

        parts.extend(body_elements)

    parts.append('</body></html>')
    return '\n'.join(parts)


# ─── PDF 转换 ──────────────────────────────────────────

def convert_pdf_to_html(file_bytes: bytes, filename: str) -> str:
    """将 PDF 文件转换为简单 HTML

    PDF 格式转换有固有限制，建议优先使用 Word (.docx) 格式导入模板。
    """
    name = filename.replace('.pdf', '').replace('.PDF', '')

    text_parts = []
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    except ImportError:
        raise RuntimeError("需要安装 PyPDF2: pip install PyPDF2")
    except Exception:
        pass

    raw_text = '\n\n'.join(text_parts).strip()

    if not raw_text:
        return _fallback_html(name, 'PDF 文件无法提取文本，请使用 Word (.docx) 格式导入模板')

    lines = raw_text.split('\n')
    html_lines = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            html_lines.append('<p>&nbsp;</p>')
            continue
        # 推断标题
        if len(stripped) < 60 and not stripped.endswith(('.', '。', '）', ')')):
            html_lines.append(f'<h2>{stripped}</h2>')
        else:
            html_lines.append(f'<p>{stripped}</p>')

    body = '\n'.join(html_lines)
    return _wrap_html(name, body)


def _wrap_html(title: str, body: str) -> str:
    return f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>{title}</title>
<style>
body{{font-family:"Microsoft YaHei","PingFang SC",sans-serif;max-width:800px;margin:0 auto;padding:20px;color:#333;line-height:1.8;}}
h1{{font-size:24px;border-bottom:2px solid #333;padding-bottom:6px;}}
h2{{font-size:18px;border-bottom:1.5px solid #333;padding-bottom:4px;margin-top:12px;}}
table{{width:100%;border-collapse:collapse;margin:8px 0;}}
td,th{{border:1px solid #ddd;padding:6px 10px;}}
</style>
</head><body>
<h1>{title}</h1>
{body}
</body></html>'''


def _fallback_html(title: str, message: str) -> str:
    return _wrap_html(title, f'<p style="color:#999;font-style:italic">({message})</p>')


def convert_to_html(file_bytes: bytes, filename: str) -> str:
    """根据文件扩展名自动选择转换器，返回 HTML 字符串"""
    ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''

    if ext in ('docx', 'doc'):
        return convert_docx_to_html(file_bytes, filename)
    elif ext == 'pdf':
        return convert_pdf_to_html(file_bytes, filename)
    else:
        raise ValueError(f"不支持的文件格式: .{ext}")
