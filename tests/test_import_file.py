# -*- coding: utf-8 -*-
"""已有简历文件导入（方案B：PDF/Word 抽文本）与采集提示词（档2）接口测试"""
import io
import os
import sys
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from fastapi.testclient import TestClient
from app import app

client = TestClient(app)


def _make_docx_bytes():
    """用 python-docx 生成带中文内容的临时 docx"""
    from docx import Document
    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("姓名：张三 | 电话：13800138000 | 邮箱：zs@example.com")
    doc.add_paragraph("教育背景：")
    doc.add_paragraph("清华大学 计算机科学与技术 硕士 2020.09-2023.06")
    doc.add_paragraph("实习经历：")
    doc.add_paragraph("字节跳动 后端开发实习生 2022.06-2022.09 负责订单系统开发")
    doc.add_paragraph("技能：")
    doc.add_paragraph("Python、SQL、Docker")
    doc.save(buf)
    return buf.getvalue()


def _make_pdf_bytes():
    """用 xhtml2pdf 生成带文本层的 PDF（英文，避免字体问题）"""
    from xhtml2pdf import pisa
    html = "<html><body><h1>Resume</h1><p>Name: Zhang San</p><p>Education: Tsinghua Univ 2020-2023</p><p>Skills: Python, SQL</p></body></html>"
    buf = io.BytesIO()
    pisa.CreatePDF(io.StringIO(html), dest=buf)
    return buf.getvalue()


def test_collect_prompt_returns_full_text():
    r = client.get("/api/experiences/collect-prompt")
    assert r.status_code == 200
    data = r.json()
    assert "prompt" in data and len(data["prompt"]) > 500
    p = data["prompt"]
    assert "简历信息采集助手" in p
    assert "绝不编造" in p
    assert "待补充" in p
    assert "work_experience" in p
    assert "others" in p


def test_import_docx_extracts_chinese_text():
    r = client.post(
        "/api/experiences/import-file",
        files={"file": ("resume.docx", _make_docx_bytes(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert r.status_code == 200, r.text
    data = r.json()
    assert data["ext"] == "docx"
    assert "清华大学" in data["text"]
    assert "字节跳动" in data["text"]
    assert "Python" in data["text"]


def test_import_pdf_extracts_text():
    r = client.post(
        "/api/experiences/import-file",
        files={"file": ("resume.pdf", _make_pdf_bytes(), "application/pdf")},
    )
    assert r.status_code == 200, r.text
    data = r.json()
    assert data["ext"] == "pdf"
    assert "Zhang San" in data["text"] or "Tsinghua" in data["text"]


def test_import_unsupported_ext_400():
    r = client.post(
        "/api/experiences/import-file",
        files={"file": ("resume.txt", b"hello world", "text/plain")},
    )
    assert r.status_code == 400


def test_import_empty_text_400():
    # 生成一个没有文字层的"空白" PDF（只有空白页）→ 应返回 400 引导信息
    from PyPDF2 import PdfWriter
    from PyPDF2.generic import DecodedStreamObject, NameObject
    writer = PdfWriter()
    page = writer.add_blank_page(width=595, height=842)
    buf = io.BytesIO()
    writer.write(buf)
    r = client.post(
        "/api/experiences/import-file",
        files={"file": ("blank.pdf", buf.getvalue(), "application/pdf")},
    )
    assert r.status_code == 400
    assert "未能从文件中提取到文字" in r.json()["detail"]
